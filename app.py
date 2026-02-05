import sys
import os
import webbrowser
import threading
import platform
import base64
import json
from io import BytesIO
from datetime import datetime
import subprocess

try:
    import pythoncom
except ImportError:
    pythoncom = None

from flask import Flask, render_template, request, jsonify, send_file, Response
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ======================== Paths & Flask Setup ========================

if getattr(sys, 'frozen', False):
    BASE_DIR = sys._MEIPASS
    STORAGE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    STORAGE_DIR = BASE_DIR

app = Flask(
    __name__,
    template_folder=os.path.join(BASE_DIR, 'templates'),
    static_folder=os.path.join(BASE_DIR, 'static')
)

DOCUMENTS_FOLDER = os.path.join(STORAGE_DIR, 'documents')
TEMPLATES_FOLDER = os.path.join(STORAGE_DIR, 'saved_templates')
SIGNATURES_FOLDER = os.path.join(STORAGE_DIR, 'signatures')

for folder in [DOCUMENTS_FOLDER, TEMPLATES_FOLDER, SIGNATURES_FOLDER]:
    if not os.path.exists(folder):
        try:
            os.makedirs(folder)
        except Exception as e:
            print(f"Error creating folder {folder}: {e}")


# ======================== Helpers ========================

def set_cell_shading(cell, color):
    """უჯრის ფონის ფერი"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_footer_text(doc, text="ფორმა 100", font_size=9):
    """
    ტექსტის დამატება გვერდის ბოლოში (Footer).
    გამოჩნდება ყველა გვერდზე (სექციაზე).
    """
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(text)
        run.font.name = 'Sylfaen'
        run.font.size = Pt(font_size)
        # თუ გინდა გამუქდეს:
        # run.bold = True


def decode_base64_image(base64_string):
    """Base64 სურათის დეკოდირება და BytesIO დაბრუნება"""
    if not base64_string or not isinstance(base64_string, str):
        return None
    if not base64_string.startswith('data:image'):
        return None
    try:
        header, data = base64_string.split(',', 1)
        image_data = base64.b64decode(data)
        return BytesIO(image_data)
    except Exception as e:
        print(f"Image decode error: {e}")
        return None


def find_libreoffice():
    if platform.system() == 'Windows':
        paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            os.path.expandvars(r"%PROGRAMFILES%\LibreOffice\program\soffice.exe"),
            os.path.expandvars(r"%PROGRAMFILES(X86)%\LibreOffice\program\soffice.exe"),
        ]
        for p in paths:
            if os.path.exists(p):
                return p
    return None


def convert_to_pdf(docx_path, output_folder):
    system = platform.system()
    pdf_path = docx_path.replace('.docx', '.pdf')

    # 1) Windows + Word (docx2pdf)
    if system == 'Windows':
        try:
            if pythoncom:
                pythoncom.CoInitialize()
            from docx2pdf import convert
            convert(docx_path, pdf_path)
            if os.path.exists(pdf_path):
                return pdf_path
        except Exception as e:
            print(f"docx2pdf failed: {e}")
        finally:
            pass

    # 2) LibreOffice (fallback)
    lo = find_libreoffice()
    if lo:
        try:
            subprocess.run([
                lo, '--headless', '--convert-to', 'pdf',
                '--outdir', output_folder, docx_path
            ], capture_output=True, timeout=60)
            if os.path.exists(pdf_path):
                return pdf_path
        except Exception as e:
            print(f"LibreOffice failed: {e}")

    return None


# ======================== Document Builders ========================

def _build_form_100_structure(data, font_size_pt):
    """
    დამხმარე ფუნქცია, რომელიც აწყობს ფორმა 100-ს.
    font_size_pt განსაზღვრავს შრიფტის ზომას (11 შენახვისთვის, 10 ბეჭდვისთვის).
    """
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)

    style = doc.styles['Normal']
    style.font.name = 'Sylfaen'
    style.font.size = Pt(font_size_pt)

    # Header
    approval = doc.add_paragraph()
    approval.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = approval.add_run(
        'დანართი №2 დამტკიცებულია საქართველოს შრომის\nჯანმრთელობისა და სოციალური დაცვის მინისტრის\n2013 წ 03.12 №01-42/ნ ბრძანებით')
    run.font.size = Pt(8)
    run.font.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(data.get('form_type', 'სამედიცინო დოკუმენტაცია ფორმა № IV-100/ა')).bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('ცნობა ჯანმრთელობის მდგომარეობის შესახებ').bold = True
    p.runs[0].font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"გაცემის თარიღი: {data.get('document_date', '')}     ბარათის №: {data.get('registration_number', '')}")

    doc.add_paragraph()

    # Tables
    # 1. Issuer
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = 'Table Grid'
    t1.rows[0].cells[0].merge(t1.rows[0].cells[1])
    t1.rows[0].cells[0].text = "1. გამცემი ორგანიზაცია"
    set_cell_shading(t1.rows[0].cells[0], "D9E2F3")
    t1.rows[1].cells[0].text = "დასახელება:"
    t1.rows[1].cells[1].text = data.get('facility_name', '')
    t1.rows[2].cells[0].text = "საიდენტიფიკაციო კოდი:"
    t1.rows[2].cells[1].text = data.get('identification_code', '')
    t1.rows[3].cells[0].text = "მისამართი:"
    t1.rows[3].cells[1].text = data.get('facility_address', '')
    doc.add_paragraph()

    # 2. Recipient
    t2 = doc.add_table(rows=2, cols=2)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].merge(t2.rows[0].cells[1])
    t2.rows[0].cells[0].text = "2. მიმღები ორგანიზაცია"
    set_cell_shading(t2.rows[0].cells[0], "D9E2F3")
    t2.rows[1].cells[0].text = "დასახელება:"
    t2.rows[1].cells[1].text = data.get('recipient_name', '')
    doc.add_paragraph()

    # 3-7. Patient
    t3 = doc.add_table(rows=6, cols=2)
    t3.style = 'Table Grid'
    t3.rows[0].cells[0].merge(t3.rows[0].cells[1])
    t3.rows[0].cells[0].text = "პაციენტის მონაცემები"
    set_cell_shading(t3.rows[0].cells[0], "E2EFDA")
    fields = [
        ("3. სახელი, გვარი:", data.get('patient_name', '')),
        ("4. დაბადების თარიღი:", data.get('birth_date', '')),
        ("5. პირადი ნომერი:", data.get('personal_id', '')),
        ("6. მისამართი:", data.get('patient_address', '')),
        ("7. სამუშაო ადგილი:", data.get('occupation', '')),
    ]
    for i, (k, v) in enumerate(fields):
        t3.rows[i + 1].cells[0].text = k
        t3.rows[i + 1].cells[1].text = str(v)
    doc.add_paragraph()

    # 8. Dates
    t4 = doc.add_table(rows=2, cols=2)
    t4.style = 'Table Grid'
    t4.rows[0].cells[0].merge(t4.rows[0].cells[1])
    t4.rows[0].cells[0].text = "8. ჰოსპიტალიზაციის ვადები"
    set_cell_shading(t4.rows[0].cells[0], "D9E2F3")
    t4.rows[1].cells[0].text = f"მიღება: {data.get('hospitalization_date', '')}"
    t4.rows[1].cells[1].text = f"გაწერა: {data.get('discharge_date', '')}"
    doc.add_paragraph()

    # 9. Diagnosis
    t5 = doc.add_table(rows=3, cols=2)
    t5.style = 'Table Grid'
    t5.rows[0].cells[0].merge(t5.rows[0].cells[1])
    t5.rows[0].cells[0].text = "9. დიაგნოზი"
    set_cell_shading(t5.rows[0].cells[0], "FCE4D6")
    t5.rows[1].cells[0].text = "ძირითადი:"
    t5.rows[1].cells[1].text = data.get('main_diagnosis', '')
    t5.rows[2].cells[0].text = "ექიმის მიერ დაზუსტება:"
    t5.rows[2].cells[1].text = data.get('case_code', '')
    doc.add_paragraph()

    # 10. Past Diseases
    t6 = doc.add_table(rows=2, cols=1)
    t6.style = 'Table Grid'
    t6.rows[0].cells[0].text = "10. გადატანილი დაავადებები"
    set_cell_shading(t6.rows[0].cells[0], "D9E2F3")
    t6.rows[1].cells[0].text = data.get('past_diseases', '')
    doc.add_paragraph()

    # 11. Anamnesis
    t7 = doc.add_table(rows=2, cols=1)
    t7.style = 'Table Grid'
    t7.rows[0].cells[0].text = "11. მოკლე ანამნეზი"
    set_cell_shading(t7.rows[0].cells[0], "D9E2F3")
    t7.rows[1].cells[0].text = data.get('anamnesis', '')
    doc.add_paragraph()

    # 12. Investigations
    t8 = doc.add_table(rows=4, cols=1)
    t8.style = 'Table Grid'
    t8.rows[0].cells[0].text = "12. ჩატარებული გამოკვლევები"
    set_cell_shading(t8.rows[0].cells[0], "D9E2F3")

    # სისხლის ანალიზი - ზომა აქედან მოდის
    cell = t8.rows[1].cells[0]
    p = cell.paragraphs[0]
    run = p.add_run(f"სისხლის საერთო ანალიზი BL.6: {data.get('blood_analysis', '')}")
    run.font.size = Pt(font_size_pt)  # <--- დინამიური ზომა

    t8.rows[2].cells[0].text = f"გლუკოზის განსაზღვრა სისხლის შრატში BL.12.1: {data.get('biochemistry', '')}"
    t8.rows[3].cells[0].text = f"ინსტრუმენტული კვლევები: {data.get('instrumental', '')}"
    doc.add_paragraph()

    # 13. Course
    t9 = doc.add_table(rows=2, cols=1)
    t9.style = 'Table Grid'
    t9.rows[0].cells[0].text = "13. დაავადების მიმდინარეობა"
    set_cell_shading(t9.rows[0].cells[0], "D9E2F3")
    course_text = f"""ტიპი: {data.get('course_type', '')}

მიღებისას: {data.get('admission_status', '')}
ვიტალური მაჩვენებლები: T-{data.get('admission_temp', '')}°C | HR-{data.get('admission_hr', '')} | BP-{data.get('admission_bp', '')} | RR-{data.get('admission_rr', '')} | SpO2-{data.get('admission_spo2', '')}

გაწერისას: {data.get('discharge_status', '')}
ვიტალური მაჩვენებლები: T-{data.get('discharge_temp', '')}°C | HR-{data.get('discharge_hr', '')} | BP-{data.get('discharge_bp', '')} | RR-{data.get('discharge_rr', '')} | SpO2-{data.get('discharge_spo2', '')}"""
    t9.rows[1].cells[0].text = course_text
    doc.add_paragraph()

    # 14. Treatment
    t10 = doc.add_table(rows=2, cols=1)
    t10.style = 'Table Grid'
    t10.rows[0].cells[0].text = "14. ჩატარებული მკურნალობა"
    set_cell_shading(t10.rows[0].cells[0], "D9E2F3")
    t10.rows[1].cells[0].text = f"მედიკამენტები:\n{data.get('medications', '')}\n\nკოდი: {data.get('treatment_code', '')}"
    doc.add_paragraph()

    # 15-17. Outcome
    t11 = doc.add_table(rows=4, cols=2)
    t11.style = 'Table Grid'
    t11.rows[0].cells[0].merge(t11.rows[0].cells[1])
    t11.rows[0].cells[0].text = "გამოსავალი"
    set_cell_shading(t11.rows[0].cells[0], "E2EFDA")
    t11.rows[1].cells[0].text = "15. სტაციონარში გადაყვანა:"
    t11.rows[1].cells[1].text = data.get('transfer_to_hospital', '-') or '-'
    t11.rows[2].cells[0].text = "16. გაწერის მდგომარეობა:"
    t11.rows[2].cells[1].text = data.get('discharge_condition', '')
    t11.rows[3].cells[0].text = "17. რეკომენდაციები:"
    t11.rows[3].cells[1].text = data.get('recommendations', '')
    doc.add_paragraph()

    # 18-20. Signatures
    t12 = doc.add_table(rows=4, cols=2)
    t12.style = 'Table Grid'
    t12.rows[0].cells[0].merge(t12.rows[0].cells[1])
    t12.rows[0].cells[0].text = "ხელმოწერები"
    set_cell_shading(t12.rows[0].cells[0], "D9E2F3")
    t12.rows[1].cells[0].text = "18. მკურნალი ექიმი:"
    t12.rows[1].cells[1].text = data.get('attending_doctor', '')
    t12.rows[2].cells[0].text = "19. დაწესებულების ხელმძღვანელი:"
    t12.rows[2].cells[1].text = data.get('facility_head', '')
    t12.rows[3].cells[0].text = "20. ცნობის გაცემის თარიღი:"
    t12.rows[3].cells[1].text = data.get('issue_date', '')
    doc.add_paragraph()

    # Electronic Signatures (Images)
    sig_table = doc.add_table(rows=2, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["ექიმის ხელმოწერა", "ბეჭედი", "ხელმძღვანელის ხელმოწერა"]
    for i, lbl in enumerate(labels):
        sig_table.rows[0].cells[i].text = lbl
        sig_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    keys = ['doctor_signature_image', 'stamp_image', 'head_signature_image']
    for i, key in enumerate(keys):
        img_data = data.get(key, '')
        img_stream = decode_base64_image(img_data)
        cell = sig_table.rows[1].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_stream:
            try:
                run = p.add_run()
                run.add_picture(img_stream, width=Inches(1.2))
            except Exception:
                p.text = "Error"
        else:
            p.text = "________________"

    return doc


def create_form_100_document_save(data):
    """შენახვისთვის: შრიფტი 11"""
    doc = _build_form_100_structure(data, font_size_pt=11)
    add_footer_text(doc, "ფორმა 100", font_size=9)  # <-- დამატებულია Footer
    return doc


def create_form_100_document_print(data):
    """ბეჭდვისთვის: შრიფტი 10"""
    doc = _build_form_100_structure(data, font_size_pt=10)
    add_footer_text(doc, "ფორმა 100", font_size=9)  # <-- დამატებულია Footer
    return doc


def create_medical_record_document(data):
    """სამედიცინო ჩანაწერი - კურსუსი"""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1)

    style = doc.styles['Normal']
    style.font.name = 'Sylfaen'
    style.font.size = Pt(11)

    # სათაური
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.add_run(data.get('facility_name', 'პრემიუმ მედ გრუპი')).bold = True
    h1.runs[0].font.size = Pt(14)
    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2.add_run(data.get('department', 'გადაუდებელი მედიცინა'))
    doc.add_paragraph()

    # პაციენტი
    t1 = doc.add_table(rows=4, cols=2)
    t1.style = 'Table Grid'
    t1.rows[0].cells[0].merge(t1.rows[0].cells[1])
    t1.rows[0].cells[0].text = "პაციენტის მონაცემები"
    set_cell_shading(t1.rows[0].cells[0], "E2EFDA")
    t1.rows[1].cells[0].text = "ბარათის №:"
    t1.rows[1].cells[1].text = data.get('card_number', '-')
    t1.rows[2].cells[0].text = "სახელი, გვარი:"
    t1.rows[2].cells[1].text = data.get('patient_name', '-')
    t1.rows[3].cells[0].text = "მიღების სტატუსი:"
    t1.rows[3].cells[1].text = data.get('admission_status', 'თვითდინებით')
    doc.add_paragraph()

    # დიაგნოზი
    t2 = doc.add_table(rows=2, cols=2)
    t2.style = 'Table Grid'
    t2.rows[0].cells[0].merge(t2.rows[0].cells[1])
    t2.rows[0].cells[0].text = "დიაგნოზი (ICD-10)"
    set_cell_shading(t2.rows[0].cells[0], "FCE4D6")
    t2.rows[1].cells[0].text = f"კოდი: {data.get('icd_code', '')}"
    t2.rows[1].cells[1].text = data.get('diagnosis_description', '')
    doc.add_paragraph()

    # ჩივილები
    t3 = doc.add_table(rows=2, cols=1)
    t3.style = 'Table Grid'
    t3.rows[0].cells[0].text = "ჩივილები"
    set_cell_shading(t3.rows[0].cells[0], "D9E2F3")
    t3.rows[1].cells[0].text = data.get('complaints', '')
    doc.add_paragraph()

    # ანამნეზი
    t4 = doc.add_table(rows=2, cols=1)
    t4.style = 'Table Grid'
    t4.rows[0].cells[0].text = "ანამნეზი"
    set_cell_shading(t4.rows[0].cells[0], "D9E2F3")
    t4.rows[1].cells[0].text = data.get('anamnesis', '')
    doc.add_paragraph()

    al = doc.add_paragraph()
    r_al = al.add_run("ალერგიები: ")
    r_al.font.bold = True
    al.add_run(data.get('allergies', 'არა'))
    doc.add_paragraph()

    # ობიექტური სტატუსი
    h_obj = doc.add_paragraph()
    h_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_obj = h_obj.add_run("ობიექტური სტატუსი")
    r_obj.font.bold = True
    r_obj.font.size = Pt(12)

    vt = doc.add_table(rows=2, cols=5)
    vt.style = 'Table Grid'
    headers = ["T°C", "BP", "HR", "RR", "SpO₂"]
    values = [
        data.get('temperature', ''),
        data.get('blood_pressure', ''),
        data.get('heart_rate', ''),
        data.get('respiratory_rate', ''),
        data.get('spo2', '')
    ]
    for i, h in enumerate(headers):
        vt.rows[0].cells[i].text = h
        set_cell_shading(vt.rows[0].cells[i], "D0D0D0")
        vt.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vt.rows[1].cells[i].text = str(values[i]) if values[i] else ''
        vt.rows[1].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    sys_t = doc.add_table(rows=9, cols=2)
    sys_t.style = 'Table Grid'
    systems = [
        ("ზოგადი მდგომარეობა:", data.get('general_condition', '')),
        ("კანი:", data.get('skin', '')),
        ("პერიფერიული შეშუპება:", data.get('edema', '')),
        ("გულ-სისხლძარღვთა:", data.get('cardiovascular', '')),
        ("სასუნთქი სისტემა:", data.get('respiratory', '')),
        ("საჭმლის მომნელებელი:", data.get('digestive', '')),
        ("შარდგამომყოფი:", data.get('urinary', '')),
        ("ნერვული სისტემა:", data.get('neurological', '')),
        ("საყრდენ-მამოძრავებელი:", data.get('musculoskeletal', '')),
    ]
    for i, (label, value) in enumerate(systems):
        sys_t.rows[i].cells[0].text = label
        sys_t.rows[i].cells[1].text = str(value) if value else ''
        set_cell_shading(sys_t.rows[i].cells[0], "F2F2F2")
    doc.add_paragraph()

    # წინასწარი დიაგნოზი
    p_pd = doc.add_paragraph()
    r_pd = p_pd.add_run("წინასწარი დიაგნოზი: ")
    r_pd.font.bold = True
    p_pd.add_run(data.get('preliminary_diagnosis', ''))

    # [1] ექიმის ხელმოწერა (დიაგნოზთან)
    p_doc = doc.add_paragraph()
    r_doc = p_doc.add_run("მკურნალი ექიმი: ")
    r_doc.font.bold = True
    p_doc.add_run(data.get('doctor', ''))

    # ხელმოწერის სურათი
    doctor_sig_data = data.get('doctor_signature_image', '')
    if doctor_sig_data:
        img_stream = decode_base64_image(doctor_sig_data)
        if img_stream:
            try:
                run = p_doc.add_run()
                run.add_picture(img_stream, width=Inches(0.8))  # მცირე ზომა ტექსტის გასწვრივ
            except:
                pass

    doc.add_page_break()

    # მიმდინარეობის ფურცელი
    h_prog = doc.add_paragraph()
    h_prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_prog = h_prog.add_run("პაციენტის მიმდინარეობის ფურცელი (დღიური)")
    r_prog.font.bold = True
    r_prog.font.size = Pt(14)
    doc.add_paragraph()

    # Initial
    t_init = doc.add_table(rows=2, cols=1)
    t_init.style = 'Table Grid'
    head_init = "პირველადი შეფასება / მიღება"
    if data.get('initial_date'):
        head_init += f"  (თარიღი: {data.get('initial_date')})"
    t_init.rows[0].cells[0].text = head_init
    set_cell_shading(t_init.rows[0].cells[0], "D9E2F3")
    t_init.rows[1].cells[0].text = data.get('initial_narrative', '')
    doc.add_paragraph()

    p_id = doc.add_paragraph()
    r_id = p_id.add_run("წინასწარი დიაგნოზი: ")
    r_id.font.bold = True
    p_id.add_run(data.get('initial_diagnosis', ''))
    doc.add_paragraph()

    # დანიშნულებები
    p_ord = doc.add_paragraph()
    r_ord = p_ord.add_run("დანიშნულებები:")
    r_ord.font.bold = True

    if data.get('investigations'):
        p_inv = doc.add_paragraph()
        r_inv = p_inv.add_run("გამოკვლევები:")
        r_inv.font.italic = True
        for line in data['investigations'].split('\n'):
            if line.strip():
                b = doc.add_paragraph(style='List Bullet')
                b.add_run(line.strip())

    if data.get('medications'):
        p_med = doc.add_paragraph()
        r_med = p_med.add_run("მედიკამენტები:")
        r_med.font.italic = True
        for line in data['medications'].split('\n'):
            if line.strip():
                b = doc.add_paragraph(style='List Bullet')
                b.add_run(line.strip())
    doc.add_paragraph()

    # [2] ექიმის ხელმოწერა (დანიშნულებებთან)
    _p_sig1 = doc.add_paragraph()
    _p_sig1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _p_sig1.add_run(f"ექიმი: {data.get('doctor_signature', '')}")

    if doctor_sig_data:
        img_stream = decode_base64_image(doctor_sig_data)
        if img_stream:
            try:
                run = _p_sig1.add_run()
                run.add_picture(img_stream, width=Inches(0.8))
            except:
                pass

    doc.add_paragraph()
    doc.add_paragraph()

    # გაწერა
    t_dis = doc.add_table(rows=2, cols=1)
    t_dis.style = 'Table Grid'
    head_dis = "გადაფასება / გაწერა"
    if data.get('discharge_note_date'):
        head_dis += f"  (თარიღი: {data.get('discharge_note_date')})"
    t_dis.rows[0].cells[0].text = head_dis
    set_cell_shading(t_dis.rows[0].cells[0], "E2EFDA")
    t_dis.rows[1].cells[0].text = data.get('discharge_narrative', '')
    doc.add_paragraph()

    # [3] ექიმის ხელმოწერა (გაწერასთან)
    _p_sig2 = doc.add_paragraph()
    _p_sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _p_sig2.add_run(f"ექიმი: {data.get('discharge_doctor', '')}")

    if doctor_sig_data:
        img_stream = decode_base64_image(doctor_sig_data)
        if img_stream:
            try:
                run = _p_sig2.add_run()
                run.add_picture(img_stream, width=Inches(0.8))
            except:
                pass

    return doc


# ======================== Routes ========================

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/api/save-document', methods=['POST'])
def save_document():
    try:
        data = request.json
        doc_type = data.get('document_type', 'form_100')
        raw_filename = data.get('filename', f'doc_{datetime.now().strftime("%H%M%S")}')
        filename = "".join(c for c in raw_filename if c.isalnum() or c in ('_', '-', ' '))

        # 1. დოკუმენტის შექმნა (DOCX)
        if doc_type == 'form_100':
            # შენახვისთვის -> დიდი შრიფტი (11)
            doc = create_form_100_document_save(data)
        else:
            doc = create_medical_record_document(data)

        docx_filename = f'{filename}.docx'
        docx_path = os.path.join(DOCUMENTS_FOLDER, docx_filename)
        doc.save(docx_path)

        # 2. PDF კონვერტაცია შენახვისას (ფორმა 100-ისთვის)
        if doc_type == 'form_100':
            pdf_path = convert_to_pdf(docx_path, DOCUMENTS_FOLDER)

            if pdf_path and os.path.exists(pdf_path):
                # PDF წარმატებით შეიქმნა
                try:
                    os.remove(docx_path)  # DOCX წაშლა (თუ გინდათ)
                except:
                    pass
                return jsonify({
                    'success': True,
                    'filename': os.path.basename(pdf_path),
                    'is_pdf': True
                })
            else:
                # PDF ვერ შეიქმნა - ვაბრუნებთ DOCX-ს
                return jsonify({
                    'success': True,
                    'filename': docx_filename,
                    'is_pdf': False,
                    'message': 'PDF ვერ შეიქმნა, ინახება DOCX'
                })
        else:
            # სხვა ტიპის დოკუმენტებისთვის (Medical Record) - ვაბრუნებთ DOCX-ს
            return jsonify({
                'success': True,
                'filename': docx_filename,
                'is_pdf': False
            })

    except Exception as e:
        print(e)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/print-document', methods=['POST'])
def print_document():
    try:
        data = request.json
        doc_type = data.get('document_type', 'form_100')

        # --- ფაილის სახელის ავტომატური გენერაცია ---

        # 1. პაციენტის სახელი (თუ ცარიელია, დაერქმევა 'Pacienti')
        patient_name = data.get('patient_name', '').strip()
        if not patient_name:
            patient_name = "Pacienti"

        # 2. თარიღი (პრიორიტეტი: დოკუმენტის თარიღი -> გაცემის თარიღი -> დღევანდელი)
        date_str = data.get('document_date') or data.get('issue_date') or datetime.now().strftime("%Y-%m-%d")

        # 3. სახელის გასუფთავება (სპეისების შეცვლა ქვედა ტირეებით და უსაფრთხო სიმბოლოები)
        # დავტოვოთ ქართული ასოები, ლათინური, ციფრები და ტირეები
        safe_chars = set(
            'აბგდევზთიკლმნოპჟრსტუფქღყშჩცძწჭხჯჰabcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789-_ ')
        clean_name = "".join(c for c in patient_name if c in safe_chars)
        clean_name = clean_name.replace(' ', '_')  # სპეისის შეცვლა ტირეთი

        # 4. საბოლოო სახელი: სახელი_გვარი + თარიღი + ფორმა_100
        if doc_type == 'form_100':
            filename = f"{clean_name}_{date_str}_ფორმა_100"
        else:
            filename = f"{clean_name}_{date_str}"

        # -------------------------------------------

        if doc_type == 'form_100':
            doc = create_form_100_document_print(data)
        else:
            doc = create_medical_record_document(data)

        docx_path = os.path.join(DOCUMENTS_FOLDER, f'{filename}.docx')
        doc.save(docx_path)

        pdf_path = convert_to_pdf(docx_path, DOCUMENTS_FOLDER)

        if pdf_path and os.path.exists(pdf_path):
            return jsonify({'success': True, 'filename': os.path.basename(pdf_path), 'is_pdf': True})
        else:
            return jsonify({'success': True, 'filename': f'{filename}.docx', 'is_pdf': False})
    except Exception as e:
        print(e)
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/print-page/<filename>')
def print_page(filename):
    filepath = os.path.join(DOCUMENTS_FOLDER, filename)
    if not os.path.exists(filepath):
        return "File not found", 404
    return render_template('print.html', filename=filename)


@app.route('/api/view-pdf/<filename>')
def view_pdf(filename):
    try:
        filepath = os.path.join(DOCUMENTS_FOLDER, filename)
        if not os.path.exists(filepath):
            return "File not found", 404

        mimetype = 'application/pdf' if filename.endswith('.pdf') else 'application/octet-stream'
        with open(filepath, 'rb') as f:
            return Response(f.read(), mimetype=mimetype)
    except Exception as e:
        return str(e), 500


@app.route('/api/download/<filename>')
def download(filename):
    return send_file(os.path.join(DOCUMENTS_FOLDER, filename), as_attachment=True)


@app.route('/api/upload-signature', methods=['POST'])
def upload_signature():
    try:
        file = request.files['file']
        sig_type = request.form.get('type', 'doctor')
        ext = file.filename.rsplit('.', 1)[-1].lower()
        filename = f'{sig_type}_signature.{ext}'
        path = os.path.join(SIGNATURES_FOLDER, filename)
        file.save(path)
        with open(path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode('utf-8')
        return jsonify({'success': True, 'base64': f'data:image/{ext};base64,{b64}'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/get-signatures')
def get_signatures():
    sigs = {}
    for t in ['doctor', 'head', 'stamp']:
        for ext in ['png', 'jpg', 'jpeg']:
            path = os.path.join(SIGNATURES_FOLDER, f'{t}_signature.{ext}')
            if os.path.exists(path):
                with open(path, 'rb') as f:
                    b64 = base64.b64encode(f.read()).decode('utf-8')
                sigs[t] = f'data:image/{ext};base64,{b64}'
                break
    return jsonify({'success': True, 'signatures': sigs})


@app.route('/api/search-patients', methods=['GET'])
def search_patients():
    query = request.args.get('q', '').lower()
    if not query:
        return jsonify({'success': True, 'results': []})

    results = []

    # 1. შენახული დოკუმენტების ძებნა (documents/)
    if os.path.exists(DOCUMENTS_FOLDER):
        for filename in os.listdir(DOCUMENTS_FOLDER):
            if filename.endswith('.docx') and query in filename.lower():
                results.append({
                    'type': 'document',
                    'name': filename,
                    'path': f'/api/download/{filename}',
                    'date': datetime.fromtimestamp(os.path.getmtime(os.path.join(DOCUMENTS_FOLDER, filename))).strftime(
                        '%Y-%m-%d %H:%M')
                })

    # 2. შაბლონების ძებნა (saved_templates/) - JSON შიგთავსში
    if os.path.exists(TEMPLATES_FOLDER):
        for filename in os.listdir(TEMPLATES_FOLDER):
            if filename.endswith('.json'):
                try:
                    path = os.path.join(TEMPLATES_FOLDER, filename)
                    with open(path, 'r', encoding='utf-8') as f:
                        data = json.load(f)

                    # ვეძებთ პაციენტის სახელს ან პირად ნომერს
                    patient_name = data.get('patient_name', '').lower()
                    personal_id = data.get('personal_id', '')
                    template_name = data.get('template_name', '').lower()

                    if query in patient_name or query in personal_id or query in template_name:
                        results.append({
                            'type': 'template',
                            'name': data.get('template_name', filename),
                            'id': filename.replace('.json', ''),
                            'patient': data.get('patient_name', '-'),
                            'date': data.get('created', '')[:16].replace('T', ' ')
                        })
                except:
                    continue

    return jsonify({'success': True, 'results': results})


@app.route('/api/templates', methods=['GET', 'POST'])
def handle_templates():
    if request.method == 'GET':
        templates = []
        if os.path.exists(TEMPLATES_FOLDER):
            for f in os.listdir(TEMPLATES_FOLDER):
                if f.endswith('.json'):
                    with open(os.path.join(TEMPLATES_FOLDER, f), 'r', encoding='utf-8') as file:
                        data = json.load(file)
                    templates.append({'id': f.replace('.json', ''), 'name': data.get('template_name', f), 'data': data})
        return jsonify({'success': True, 'templates': templates})

    if request.method == 'POST':
        data = request.json
        name = data.get('template_name', 'Template')
        fname = f"{name.replace(' ', '_')}_{datetime.now().strftime('%H%M%S')}.json"
        with open(os.path.join(TEMPLATES_FOLDER, fname), 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)
        return jsonify({'success': True})


@app.route('/api/templates/<tid>', methods=['DELETE'])
def delete_template(tid):
    for f in os.listdir(TEMPLATES_FOLDER):
        if f.startswith(tid):
            os.remove(os.path.join(TEMPLATES_FOLDER, f))
            return jsonify({'success': True})
    return jsonify({'success': False}), 404


# ======================== Main ========================

def open_browser():
    webbrowser.open('http://127.0.0.1:5000')


if __name__ == '__main__':
    print("=" * 50)
    print("🏥 სამედიცინო დოკუმენტაცია")
    print("=" * 50)

    lo = find_libreoffice()
    if lo:
        print(f"✅ LibreOffice: {lo}")
    else:
        print("⚠️  LibreOffice ვერ მოიძებნა")

    print(f"\n📁 დოკუმენტები: {DOCUMENTS_FOLDER}")
    print(f"📁 შაბლონები: {TEMPLATES_FOLDER}")
    print(f"📁 ხელმოწერები: {SIGNATURES_FOLDER}")
    print("\n🌐 მისამართი: http://127.0.0.1:5000")
    print("=" * 50)

    threading.Timer(1.5, open_browser).start()
    app.run(host='127.0.0.1', port=5000)